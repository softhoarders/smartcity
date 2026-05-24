#!/usr/bin/env python3
"""Generate PRESENTATION_2MIN_RO.docx (technical demo script)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT_PATHS = [
    ROOT / "PRESENTATION_2MIN_RO.docx",
    Path("/Users/thechallenger_/Downloads/pitch_technical_ro.docx"),
]


def set_cell_shading(cell, fill: str) -> None:
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {qn("w:fill"): fill, qn("w:val"): "clear"})
    shading.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], "E8EEF7")
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = val
    doc.add_paragraph()


def add_step(doc: Document, title: str, blocks: list[tuple[str, str]]) -> None:
    doc.add_heading(title, level=2)
    for label, text in blocks:
        p = doc.add_paragraph()
        p.add_run(f"{label}: ").bold = True
        p.add_run(text)


def main() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Spotflow — Demo tehnic (2 minute)", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    meta.add_run("Public: ").bold = True
    meta.add_run("ingineri, integratori, operatori IT\n")
    meta.add_run("Focus: ").bold = True
    meta.add_run("pipeline edge → server → UI; sincronizare număr la rezervare; ledger Credits; routing și reputație\n")
    meta.add_run("Durată: ").bold = True
    meta.add_run("~120 secunde")

    doc.add_heading("Pregătire (off-camera)", level=1)
    add_table(
        doc,
        ["Item", "Detaliu"],
        [
            ["Server", "python app.py în server/ (port din env, de ob. 2026)"],
            ["Viewport", "Desktop ≥1200px — hărți Leaflet, fără coliziune cu bara mobilă"],
            ["Intrare", "/login?demo=1 → Șofer → cod afișat pe pagina 2FA"],
            ["Sesiune", "Date sintetice București / Cluj / Craiova; Credits preîncărcat"],
            ["Evită", "Buton hero Explore demo; panoul concierge AI"],
        ],
    )

    doc.add_heading("Beat sheet tehnic", level=1)
    add_table(
        doc,
        ["Timp", "Strat", "Mesaj cheie"],
        [
            ["0:00–0:12", "Stack", "Edge OCR + Flask + portal — model de date unificat"],
            ["0:12–0:32", "Incident", "Heartbeat → alertă detected vs expected; lanț dovadă"],
            ["0:32–0:52", "Discovery", "Geocoding + availability + routing direct vs walk-in"],
            ["0:52–1:12", "Tranzacție", "Booking → debit wallet → număr autorizat pe device"],
            ["1:12–1:32", "Policy", "Trust passport, min_trust, smart pricing floor/ceiling"],
            ["1:32–2:00", "Close", "Driver / owner / admin; API REST extensibil"],
        ],
    )

    doc.add_heading("Script — pas cu pas", level=1)
    tip = doc.add_paragraph()
    tip.add_run("Notă: ").bold = True
    tip.add_run("Vorbește peste click-uri. Pauze scurte la 2FA și după confirmarea plății.")

    add_step(
        doc,
        "0:00 — Autentificare și context sesiune (12 s)",
        [
            ("Ecran", "/login?demo=1 → 2FA → /portal"),
            ("Acțiuni", "Autentificare șofer; cod din ecran; Continuă"),
            (
                "Spune",
                "Spotflow e un strat central Flask peste SQLite: device-uri edge raportează starea locului, "
                "iar același backend servește portalul șofer, marketplace-ul și dashboard-ul operator. "
                "Intru într-o sesiune demo — date sintetice, dar fluxurile reale.",
            ),
            ("Arată", "Număr verificat, sold Credits, sumar alerte"),
        ],
    )

    add_step(
        doc,
        "0:12 — Pipeline alertă și dovadă (20 s)",
        [
            ("Ecran", "/portal — hartă + listă alerte"),
            ("Acțiuni", "Hartă Where alerts happened; alertă Mismatch; Detected vs Expected"),
            (
                "Spune",
                "Fiecare incident e legat de un Device cu coordonate. Edge trimite heartbeat și, "
                "la nepotrivire, un Fine cu timestamp și scor OCR. Request photo pune device-ul în coadă "
                "de captură; contestația poate trece prin review automat, apoi escaladare admin.",
            ),
            ("Tehnic", "SSE pe /stream actualizează dashboard-ul fără refresh complet"),
        ],
    )

    add_step(
        doc,
        "0:32 — Discovery: hartă, filtre, routing (20 s)",
        [
            ("Ecran", "/portal/find-parking"),
            ("Acțiuni", "Find parking → Piata Universitatii → search → legendă hartă → listare Available now"),
            (
                "Spune",
                "Geocoding ancorează harta; enrich_listing_items calculează disponibilitatea la target_at. "
                "Routing separă locuri la destinație de variante walk-in — două strategii, nu doar distanță.",
            ),
            ("Nu folosi", "Panoul Quick parking search / concierge"),
        ],
    )

    add_step(
        doc,
        "0:52 — Rezervare și ledger (20 s)",
        [
            ("Acțiuni", "Pay & park → confirmă → Done"),
            (
                "Spune",
                "Credits stocate în hundredths. La submit: debit wallet + booking. Manual approval = plată "
                "rezervată; auto-approve = număr chiriaș devine assigned_plate temporar pe device, fără SSH pe Pi.",
            ),
            ("Ecran", "Payment received — pending approval sau confirmare instant — ambele valide"),
        ],
    )

    add_step(
        doc,
        "1:12 — Policy layer: trust și owner (20 s)",
        [
            ("Ecran", "/portal/settings → My spots"),
            (
                "Spune",
                "Trust passport + min_trust_score controlează rezervarea instant. Owner: approval manual/auto, "
                "smart pricing floor/ceiling. Activity log alimentează semnale de cerere pentru tarife.",
            ),
        ],
    )

    add_step(
        doc,
        "1:32 — Închidere (28 s)",
        [
            (
                "Spune",
                "Edge raportează, serverul persistă și notifică, portalul leagă monitorizarea de marketplace. "
                "Operator: /admin — flotă, cozi, export CSV. API REST pentru register, heartbeat, config, fines. "
                "Întrebări pe fluxul de date?",
            ),
            ("Opțional", "Admin → hartă flotă + chart 7 zile + appeals"),
        ],
    )

    doc.add_heading("Teleprompter", level=1)
    tele = (
        "Spotflow e un strat central Flask peste SQLite: device-uri edge raportează starea locului, "
        "iar același backend servește portalul șofer, marketplace-ul și dashboard-ul operator. "
        "Intru într-o sesiune demo — date sintetice, dar fluxurile reale.\n\n"
        "Fiecare incident e legat de un device cu coordonate. Edge trimite heartbeat și fine cu scor OCR. "
        "Request photo → coadă captură; contestație → review automat → admin. SSE actualizează dashboard-ul.\n\n"
        "Find parking: geocoding, enrichment disponibilitate, routing direct vs walk-in. "
        "Credits în hundredths; booking debitează wallet și propagă numărul autorizat când approval mode permite.\n\n"
        "Trust passport, min_trust_score, smart pricing owner. Admin: aceeași flotă, export, cozi. "
        "API REST — integrare fără UI. Întrebări?"
    )
    tp = doc.add_paragraph(tele)
    for run in tp.runs:
        run.italic = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    doc.add_heading("Checklist live", level=1)
    for item in [
        "Desktop width, server pornit",
        "/login?demo=1 → 2FA → /portal",
        "Hartă alerte + Mismatch (Detected vs Expected)",
        "Find parking → Piata Universitatii → Pay & park",
        "Explică pending vs auto-approve → Done",
        "My spots → trust / smart pricing",
        "Închidere stack + API",
    ]:
        doc.add_paragraph(f"☐ {item}", style="List Bullet")

    doc.add_heading("Recuperare rapidă", level=1)
    add_table(
        doc,
        ["Simptom", "Acțiune"],
        [
            ["Hartă goală", "Refresh; re-search Bucharest city center"],
            ["Fără Pay & park", "All nearby spots → Book anyway"],
            ["Butoane acoperite", "Lățime desktop sau scroll deasupra barei"],
            ["Concierge error", "Ignoră — search clasic"],
            ["Pending approval", "Explică manual approval; sync camera după Accept"],
        ],
    )

    doc.add_heading("Rute utile", level=1)
    add_table(
        doc,
        ["Ecran", "Cale"],
        [
            ["Login demo", "/login?demo=1"],
            ["Portal șofer", "/portal"],
            ["Marketplace", "/portal/find-parking"],
            ["Cont / trust", "/portal/settings"],
            ["Owner", "/portal/my-spots"],
            ["Admin", "/admin"],
        ],
    )

    foot = doc.add_paragraph()
    foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = foot.add_run("Verificat live — concierge exclus din traseul recomandat.")
    r.italic = True
    r.font.size = Pt(9)

    for out in OUT_PATHS:
        doc.save(out)
        print(f"Wrote {out}")


if __name__ == "__main__":
    main()
